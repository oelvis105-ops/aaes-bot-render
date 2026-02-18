# utils.py  –  clean version
import os
import io
import tempfile
import logging
from pathlib import Path
import aiohttp
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from docx import Document
import PyPDF2
from groq import Groq
from google.oauth2 import service_account
from telegram import InlineKeyboardButton, InlineKeyboardMarkup,Update, InputFile
from telegram.ext import ContextTypes
client = Groq(api_key=os.getenv("GROQ_API_KEY"))
logger = logging.getLogger("aaes-bot")

SERVICE_ACCOUNT_FILE = os.getenv("GOOGLE_SA_PATH", "service_account.json")
SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]
drive_service = build("drive", "v3", credentials=service_account.Credentials.from_service_account_file(SERVICE_ACCOUNT_FILE, scopes=SCOPES))

# ---------- Drive wrapper for AI ----------
from drive_search import search_drive as _folder_search   # NEW folder-first search

def search_drive(keyword: str) -> str:
    """
    Wraps the folder-first search so the AI module keeps working.
    Returns a plain-text list of files (or 'No matching slide found.').
    """
    result = _folder_search(keyword, mode="all", level=None)
    return result if result else "No matching slide found."

# ---------- File readers ----------
def read_file_content(file_path: str) -> str:
    try:
        if file_path.lower().endswith(".pdf"):
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() or ""
                return text.strip()
        elif file_path.lower().endswith(".docx"):
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text.strip()
    except Exception as e:
        logger.error(f"Error reading file content: {e}")
        return "Failed to read file content."

# ---------- AI chat ----------
async def ai_chat_response(prompt: str) -> str:
    try:
        # Search for relevant information in Google Drive
        drive_hint = search_drive(prompt)

        # Define the system message for the AI
        system_msg = (
            "You are AeroTutor, the AAES academic assistant. "
            "Explain things with clarity. Use engineering examples. "
            "If Google Drive reference is available, use it."
        )
        messages = [{"role": "system", "content": system_msg}]

        # Add the drive hint if available
        if drive_hint and isinstance(drive_hint, str) and "No matching" not in drive_hint:
            messages.append({"role": "system", "content": str(drive_hint)})

        # Add the user's prompt
        messages.append({"role": "user", "content": str(prompt)})

        # Generate the response from the AI
        resp = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=messages,
            temperature=0.4,
            max_tokens=400
        )
        return resp.choices[0].message.content.strip()

    except Exception as e:
        logger.error("Groq API error: %s", e)
        return (
            "⚠️ I couldn’t reach the AI service right now.\n"
            "Please try again in a few seconds or rephrase your question."
        )

async def read_file(doc, update) -> str:
    """
    Download the Telegram document and return its text content.
    Supports PDF and DOCX; returns empty string for other types.
    """
    # 1. Download the file to a temporary location
    tg_file = await doc.get_file()
    suffix = Path(doc.file_name or "file").suffix.lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp_path = tmp.name
    await tg_file.download_to_drive(tmp_path)

    # 2. Extract text
    text = ""
    try:
        if suffix == ".pdf":
            with open(tmp_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() or ""
        elif suffix == ".docx":
            document = Document(tmp_path)
            for para in document.paragraphs:
                text += para.text + "\n"
    except Exception as e:
        logger.warning("Could not extract text from file: %s", e)
    finally:
        # 3. Clean up
        try:
            os.remove(tmp_path)
        except Exception:
            pass

    return text.strip()
def prettify_answer(raw: str) -> str:
    """
    Converts a plain-text AI answer into a Telegram-friendly Markdown block.
    """
    lines = raw.splitlines()
    out = ["📘 *Quick Answer*"]
    in_list = False

    for ln in lines:
        ln = ln.strip()
        if not ln:
            continue

        # Headings like "**How it works:**"
        if ln.startswith("**") and ln.endswith("**"):
            out.append(f"\n🔹 *{ln[2:-2]}*")
            continue

        # Bullet lists
        if ln.startswith("* "):
            out.append(f"  • {ln[2:]}")
            continue

        # Numbered steps
        if ln[0].isdigit() and ln[1] == ".":
            out.append(f"  {ln}")
            continue

        # Regular paragraph
        out.append(ln)

    # Compact footer
    out.append("\n💡 _Need deeper detail? Ask follow-up or upload slides._")
    return "\n".join(out)

# ---------- File upload handler ----------
async def handle_upload(doc, update):
    try:
        file = await doc.get_file()
        path = f"uploads/{doc.file_name}"
        os.makedirs("uploads", exist_ok=True)
        await file.download_to_drive(path)

        extracted = read_file(path)
        if not extracted:
            await update.message.reply_text("I could not read this file.")
            return
        await update.message.reply_text(extracted[:2000])
        return extracted
    except Exception as e:
        logging.error(f"Upload error: {e}")
        await update.message.reply_text("Upload failed.")

def read_file_content(file_id: str) -> str:
    try:
        request = drive_service.files().get_media(fileId=file_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
        fh.seek(0)
        if file_id.endswith(".pdf"):
            reader = PyPDF2.PdfReader(fh)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text.strip()
        elif file_id.endswith(".docx"):
            doc = Document(fh)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text.strip()
    except Exception as e:
        logger.error(f"Error reading file content: {e}")
    return "Failed to read file content."

def list_files_in_folder(folder_id: str) -> list:
    query = f"'{folder_id}' in parents and trashed = false"
    results = drive_service.files().list(q=query, fields="files(id, name)").execute()
    return results.get("files", [])

def create_buttons_from_files(file_ids: list) -> InlineKeyboardMarkup:
    keyboard = []
    for file in file_ids:
        button = InlineKeyboardButton(file["name"], callback_data=f"exec_file:{file['id']}")
        keyboard.append([button])
    return InlineKeyboardMarkup(keyboard)

async def display_exec_info(update: Update, context: ContextTypes.DEFAULT_TYPE):
    q = update.callback_query
    await q.answer()
    code = q.data
    file_id = code.split(":", 1)[1]

    # Extract picture and text from the file
    picture_path, text = extract_picture_and_text(file_id)

    # Send the picture with text as a caption
    with open(picture_path, "rb") as photo:
        await context.bot.send_photo(
            chat_id=q.message.chat_id,
            photo=photo,
            caption=text,
            parse_mode="Markdown"
        )

def extract_picture_and_text(file_id: str) -> (str, str):
    try:
        request = drive_service.files().get_media(fileId=file_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()
        fh.seek(0)
        if file_id.endswith(".pdf"):
            reader = PyPDF2.PdfReader(fh)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            # Assuming the first image in the PDF is the one we want
            image = reader.pages[0].images[0]
            image_data = io.BytesIO(image.data)
            picture_path = f"{os.path.splitext(file_id)[0]}.png"
            with open(picture_path, "wb") as img_file:
                img_file.write(image_data.getvalue())
            return picture_path, text.strip()
    except Exception as e:
        logger.error(f"Error extracting picture and text: {e}")
    return "default_image.png", "Failed to extract picture and text."

# ---------- NEW: send actual file to Telegram (up to 1.9 GB) ----------
from google.oauth2 import service_account
from googleapiclient.discovery import build
SERVICE_ACCOUNT_FILE = os.getenv("GOOGLE_SA_PATH", "service_account.json")
SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]   # removed trailing space
def _drive_service():
    creds = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE, scopes=SCOPES
    )
    return build("drive", "v3", credentials=creds)

import mimetypes
from telegram.constants import ParseMode

async def send_drive_file(file_id: str,
                          filename: str,
                          bot,
                          chat_id: int,
                          caption: str | None = None,
                          thumb_path: str | None = None):
    """
    Download anything from Drive and send it to Telegram.
    Supports ZIP, RAR, MP4, PDF, DOCX, etc.
    Falls back to a Drive link if the file is too big or Telegram rejects it.
    """
    try:
        drive = _drive_service()
        request = drive.files().get_media(fileId=file_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request, chunksize=10 * 1024 * 1024)
        done = False
        while not done:
            _, done = downloader.next_chunk()
        fh.seek(0)

        # Guess MIME type
        mime, _ = mimetypes.guess_type(filename)
        if not mime:
            mime = "application/octet-stream"

        # Telegram bot limit is 50 MB for send_document
        file_size = fh.getbuffer().nbytes
        if file_size > 49 * 1024 * 1024:          # 49 MB safety margin
            file = drive.files().get(fileId=file_id, fields="webViewLink").execute()
            link = file.get("webViewLink", "https://drive.google.com")
            await bot.send_message(
                chat_id=chat_id,
                text=f"📁 *{filename}* is too large for Telegram.\n\n[Download here]({link})",
                parse_mode=ParseMode.MARKDOWN,
                disable_web_page_preview=True
            )
            return

        # Send the actual file
        await bot.send_document(
            chat_id=chat_id,
            document=InputFile(fh, filename=filename),
            caption=caption or "",
            parse_mode=ParseMode.MARKDOWN,
            read_timeout=300,
            write_timeout=300,
            thumb=open(thumb_path, "rb") if thumb_path and os.path.exists(thumb_path) else None
        )

    except Exception as e:
        logger.error("send_drive_file failed: %s", e)
        # Fallback to Drive link
        try:
            drive = _drive_service()
            file = drive.files().get(fileId=file_id, fields="webViewLink").execute()
            link = file.get("webViewLink", "https://drive.google.com")
            await bot.send_message(
                chat_id=chat_id,
                text=f"⚠️ Could not send *{filename}* directly.\n\n[Download here]({link})",
                parse_mode=ParseMode.MARKDOWN
            )
        except Exception as e2:
            logger.error("Fallback link also failed: %s", e2)
            await bot.send_message(chat_id=chat_id, text="❌ File unavailable.")

from googleapiclient.discovery import build

async def send_drive_file(file_id, filename, bot, chat_id, caption="", thumb_path=None):
    try:
        # Set up Google Drive service account
        creds = service_account.Credentials.from_service_account_file(
            'credentials.json', scopes=['https://www.googleapis.com/auth/drive.readonly']
        )
        drive_service = build('drive', 'v3', credentials=creds)

        # Download the file from Google Drive
        request = drive_service.files().get_media(fileId=file_id)
        response = request.execute()
        file_data = response.get('downloadUrl')

        # Download the file content
        async with aiohttp.ClientSession() as session:
            async with session.get(file_data, timeout=60) as resp:
                file_content = await resp.read()

        # Send the file to Telegram
        if thumb_path and os.path.exists(thumb_path):
            thumb = InputFile(thumb_path)
        else:
            thumb = None

        await bot.send_document(
            chat_id=chat_id,
            document=InputFile(io.BytesIO(file_content), filename=filename),
            thumb=thumb,
            caption=caption or "",
            parse_mode='Markdown',
            read_timeout=300,
            write_timeout=300,
        )

    except Exception as e:
        logger.error(f"Failed to send file: {e}")
        await bot.send_message(chat_id=chat_id, text="Failed to send file.")
from PIL import Image

def extract_picture_and_text(file_path: str) -> (str, str):
    try:
        if file_path.lower().endswith(".pdf"):
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() or ""
                # Assuming the first image in the PDF is the one we want
                image = reader.pages[0].images[0]
                image_data = io.BytesIO(image.data)
                picture_path = f"{os.path.splitext(file_path)[0]}.png"
                with open(picture_path, "wb") as img_file:
                    img_file.write(image_data.getvalue())
                return picture_path, text.strip()
    except Exception as e:
        logger.error(f"Error extracting picture and text: {e}")
    return "default_image.png", "Failed to extract picture and text."
from telegram import InlineKeyboardButton, InlineKeyboardMarkup

def create_buttons_from_files(file_paths: list) -> InlineKeyboardMarkup:
    keyboard = []
    for file_path in file_paths:
        file_name = os.path.basename(file_path)
        button = InlineKeyboardButton(file_name, callback_data=f"exec_file:{file_name}")
        keyboard.append([button])
    return InlineKeyboardMarkup(keyboard)

async def display_exec_info(update: Update, context: ContextTypes.DEFAULT_TYPE):
    q = update.callback_query
    await q.answer()
    code = q.data
    file_name = code.split(":", 1)[1]

    # Assuming the files are stored in a specific directory
    file_path = os.path.join("exec_files", file_name)

    # Extract picture and text from the file
    picture_path, text = extract_picture_and_text(file_path)

    # Send the picture with text as a caption
    with open(picture_path, "rb") as photo:
        await context.bot.send_photo(
            chat_id=q.message.chat_id,
            photo=photo,
            caption=text,
            parse_mode="Markdown"
        )

# ---------- conversation helpers ----------
from collections import deque

CONVERSATION = {}
MAX_HISTORY = 10

def add_turn(uid: int, role: str, text: str, file_text: str | None = None):
    """Add a message to the user’s rolling buffer."""
    if uid not in CONVERSATION:
        CONVERSATION[uid] = deque(maxlen=MAX_HISTORY)
    CONVERSATION[uid].append({"role": role, "text": text, "file": file_text})

def build_prompt(uid: int) -> str:
    """Assemble last 10 turns into a single prompt for the LLM."""
    buf = CONVERSATION.get(uid, [])
    lines = []
    for m in buf:
        if m["file"]:
            lines.append(f"User (with file): {m['text']}\nFile content:\n{m['file'][:3000]}")
        else:
            lines.append(f"{m['role'].capitalize()}: {m['text']}")
    return "\n\n".join(lines)
# ---------- END OF utils.py ----------